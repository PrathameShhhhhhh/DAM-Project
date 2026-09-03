from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Dam Monitoring System - Workflow and Implementation Plan.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level <= 2 else RGBColor(31, 77, 120)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "F2F4F7")
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 3" else 4)


def main():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Dam Monitoring System: Workflow and Implementation Plan")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("Summary of recommended improvements, technology stack, workflow, implementation phases, and project file system layout.")
    sr.font.name = "Calibri"
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "Project Goal", 1)
    doc.add_paragraph(
        "The project should evolve from a frontend-only simulation into a complete software-based dam monitoring system. "
        "It will use simulated IoT sensor readings, store them in a database, run AI/ML prediction logic, expose results through a backend API, "
        "and display risk status, history, alerts, and emergency controls on the dashboard."
    )

    add_heading(doc, "Current State", 1)
    add_bullets(doc, [
        "Existing frontend files: index.html, css/style.css, and js/app.js.",
        "The dashboard already simulates water level, rainfall, rise rate, risk score, time-to-critical, charts, CSV export, siren, SMS log, and threshold settings.",
        "There is no Flask backend, MySQL database, Python ML model, dataset folder, or API connection yet.",
    ])

    add_heading(doc, "Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Language / Technology", "Purpose"],
        [
            ["Frontend", "HTML, CSS, JavaScript", "Build the dashboard, controls, tabs, alert views, and browser interactions."],
            ["Charts", "Chart.js", "Render live telemetry and historical analytics graphs."],
            ["Backend", "Python + Flask", "Create APIs that connect frontend, database, and ML prediction logic."],
            ["Database", "MySQL", "Store sensor readings, predictions, alerts, users, and threshold settings."],
            ["AI/ML", "Python, Pandas, NumPy, Scikit-learn", "Clean data, train the risk prediction model, and produce live predictions."],
            ["Data Format", "CSV / JSON", "Represent simulated IoT readings and API request/response payloads."],
            ["Alerts", "JavaScript + backend alert APIs; optional SMS service", "Trigger dashboard warnings, siren simulation, and SMS-style notifications."],
        ],
        [1500, 2500, 5360],
    )

    add_heading(doc, "System Workflow", 1)
    add_numbered(doc, [
        "Generate simulated IoT readings for water level, rainfall, and rate of rise.",
        "Send each reading to the Flask backend through an API endpoint.",
        "Validate the incoming reading and store it in MySQL.",
        "Pass the latest sensor values to the trained ML prediction function.",
        "Classify risk as Safe, Warning, High Risk, or Critical and calculate probability plus time-to-critical.",
        "Store prediction results and any alert events in MySQL.",
        "Return current readings, predictions, history, and alerts to the frontend dashboard.",
        "Display live status, graphs, audit logs, threshold settings, siren state, and emergency notifications.",
    ])

    add_heading(doc, "Implementation Plan", 1)
    add_table(
        doc,
        ["Phase", "Owner / Module", "Implementation Work"],
        [
            ["1", "Data / IoT Simulation", "Create sample dataset files and a Python generator for realistic live readings."],
            ["2", "AI / ML", "Train a Random Forest or Decision Tree model using water level, rainfall, and rise rate."],
            ["3", "Backend + Database", "Create Flask APIs, MySQL tables, configuration, and database helper functions."],
            ["4", "ML + Backend Integration", "Load model.pkl in Flask and expose prediction endpoints."],
            ["5", "Frontend Integration", "Replace browser-only data with API calls for current data, history, predictions, and alerts."],
            ["6", "Alert System", "Store alert history, trigger critical warnings, and optionally integrate SMS provider APIs."],
            ["7", "Testing + Documentation", "Test full workflow and add README, setup steps, sample data, and validation notes."],
        ],
        [900, 2100, 6360],
    )

    add_heading(doc, "Recommended Improvements", 1)
    add_bullets(doc, [
        "Add backend APIs for current data, history, prediction, alerts, and threshold updates.",
        "Add persistent MySQL storage instead of keeping history only in browser memory.",
        "Add a real ML model and save it as model.pkl for backend inference.",
        "Add validation for sensor readings, thresholds, and emergency-control inputs.",
        "Add login/admin roles so only authorized users can modify thresholds or emergency actions.",
        "Add persistent alert history and optional SMS provider integration.",
        "Improve mobile responsiveness, accessibility, and dashboard usability.",
        "Add tests for prediction logic, API routes, database operations, and alert triggering.",
        "Add README.md, requirements.txt, .env.example, and setup instructions.",
    ])

    add_heading(doc, "Recommended Project File System Layout", 1)
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.space_after = Pt(10)
    cr = code.add_run(
        "DAM Project/\n"
        "  index.html\n"
        "  css/\n"
        "    style.css\n"
        "  js/\n"
        "    app.js\n"
        "    api.js\n"
        "  data/\n"
        "    dam_data.csv\n"
        "    generate_data.py\n"
        "    preprocess.py\n"
        "  ml/\n"
        "    train_model.py\n"
        "    predict.py\n"
        "    evaluate.py\n"
        "    model.pkl\n"
        "  backend/\n"
        "    app.py\n"
        "    config.py\n"
        "    database.py\n"
        "    requirements.txt\n"
        "    routes/\n"
        "      sensor.py\n"
        "      prediction.py\n"
        "      alerts.py\n"
        "      settings.py\n"
        "  database/\n"
        "    schema.sql\n"
        "    seed.sql\n"
        "  docs/\n"
        "    workflow-and-implementation-plan.docx\n"
        "  .env.example\n"
        "  README.md\n"
    )
    cr.font.name = "Courier New"
    cr.font.size = Pt(9)

    add_heading(doc, "Database Tables", 1)
    add_table(
        doc,
        ["Table", "Main Fields", "Purpose"],
        [
            ["sensor_data", "id, timestamp, water_level, rainfall, rise_rate", "Stores incoming sensor or simulated IoT readings."],
            ["predictions", "id, sensor_data_id, risk_level, probability, time_to_critical, timestamp", "Stores AI/ML prediction results."],
            ["alerts", "id, risk_level, alert_type, message, zone, timestamp, status", "Stores warning, high-risk, and critical alert events."],
            ["users", "id, name, phone, role, zone", "Stores emergency contacts, authorities, and possible admin users."],
            ["thresholds", "id, safe_max, warning_max, high_max, critical_max, updated_at", "Stores configurable risk thresholds."],
        ],
        [1600, 3900, 3860],
    )

    add_heading(doc, "Suggested API Endpoints", 1)
    add_bullets(doc, [
        "GET /api/current-data - latest sensor reading and prediction.",
        "GET /api/history - historical sensor and prediction records.",
        "POST /api/sensor-data - receive a new simulated IoT reading.",
        "GET /api/prediction - calculate or fetch prediction for latest data.",
        "GET /api/alerts - return stored alerts.",
        "POST /api/alerts - create or trigger an alert.",
        "POST /api/update-threshold - update warning, high-risk, and critical thresholds.",
    ])

    add_heading(doc, "Recommended Build Order", 1)
    add_numbered(doc, [
        "Create the dataset and generator first.",
        "Train and test the ML model with the dataset.",
        "Build the Flask backend and MySQL schema.",
        "Connect the ML model to the backend.",
        "Connect the frontend dashboard to backend APIs.",
        "Add alert persistence and optional SMS integration.",
        "Test the complete flow from data generation to dashboard alert.",
    ])

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Dam Monitoring System Implementation Plan")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(100, 116, 139)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
